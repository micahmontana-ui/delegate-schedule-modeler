"""
Module 7b: Word Report Builder

Pure function interface:
    build_report(output_path, summary_stats, assumptions, non_attendance,
                 stop_summary, bus_trips_df, required_activities, changelog)
        -> path (str)
"""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any

import pandas as pd


def build_report(
    output_path: str,
    summary_stats: dict[str, Any],
    assumptions: list[str],
    non_attendance: pd.DataFrame,
    stop_summary: dict[str, dict],
    bus_trips_df: pd.DataFrame,
    required_activities: dict[str, dict],
    changelog: list[str] | None = None,
) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document()

    # Title
    title = doc.add_heading("Convention Logistics Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {date.today()}")
    doc.add_paragraph("")

    # ----------------------------------------------------------------
    # Executive Summary
    # ----------------------------------------------------------------
    doc.add_heading("Executive Summary", 1)

    n_groups = summary_stats.get("total_groups", "N/A")
    n_delegates = summary_stats.get("total_delegates", "N/A")
    n_trips = summary_stats.get("total_trips", "N/A")
    n_under = summary_stats.get("under_min_trips", "N/A")
    pct_under = summary_stats.get("pct_under_min", "N/A")

    exec_para = doc.add_paragraph()
    exec_para.add_run(
        f"This model covers {n_delegates} delegates in {n_groups} groups. "
        f"The bus assignment engine produced {n_trips} total trips, of which "
        f"{n_under} ({pct_under}%) are below the minimum passenger threshold. "
    )

    # Biggest findings
    biggest = summary_stats.get("biggest_findings", [])
    if biggest:
        doc.add_paragraph("Key findings:")
        for finding in biggest:
            doc.add_paragraph(finding, style="List Bullet")

    # ----------------------------------------------------------------
    # Population Overview
    # ----------------------------------------------------------------
    doc.add_heading("Population Overview", 1)
    doc.add_paragraph(
        f"Total delegate count: {n_delegates} (see 'Delegate Groups' sheet). "
        f"Total groups: {n_groups}."
    )
    if summary_stats.get("clamped_groups"):
        doc.add_paragraph(
            f"Note: {summary_stats['clamped_groups']} groups had check-in clamped to the "
            "window start date due to their stay length exceeding the window."
        )

    # ----------------------------------------------------------------
    # Activity Scheduling
    # ----------------------------------------------------------------
    doc.add_heading("Activity Scheduling Results", 1)
    doc.add_paragraph(
        "The following table shows assignment rates for each required activity. "
        "See the 'Non-Attendance Tracking' sheet for individual non-attendance records."
    )

    act_table = doc.add_table(rows=1, cols=4)
    act_table.style = "Table Grid"
    hdr_cells = act_table.rows[0].cells
    for i, h in enumerate(["Activity", "Assigned", "Not Scheduled", "Non-Attendance Reasons"]):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    for act_name in required_activities:
        row_cells = act_table.add_row().cells
        if not non_attendance.empty:
            n_not = len(non_attendance[non_attendance["Activity"] == act_name])
        else:
            n_not = 0
        n_assigned = summary_stats.get("activity_assigned", {}).get(act_name, "N/A")
        reasons = ""
        if not non_attendance.empty:
            act_na = non_attendance[non_attendance["Activity"] == act_name]
            if not act_na.empty:
                reason_counts = act_na["Reason"].str.split(":").str[0].value_counts()
                reasons = "; ".join(f"{r}: {c}" for r, c in reason_counts.items())
        row_cells[0].text = act_name
        row_cells[1].text = str(n_assigned)
        row_cells[2].text = str(n_not)
        row_cells[3].text = reasons

    # ----------------------------------------------------------------
    # Bus Assignment Summary
    # ----------------------------------------------------------------
    doc.add_heading("Bus Assignment Summary", 1)
    doc.add_paragraph(
        f"Total bus trips: {n_trips}. Under-minimum trips: {n_under} ({pct_under}%). "
        "See 'Bus Assignments' sheet for full trip detail."
    )

    # Stop-level breakdown
    if stop_summary:
        doc.add_paragraph("By stop:")
        stop_table = doc.add_table(rows=1, cols=4)
        stop_table.style = "Table Grid"
        hdr = stop_table.rows[0].cells
        for i, h in enumerate(["Stop", "Total Trips", "Under-Min Trips", "Total Passengers"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].font.bold = True
        for stop, stats in sorted(stop_summary.items()):
            rc = stop_table.add_row().cells
            rc[0].text = stop
            rc[1].text = str(stats["total_trips"])
            rc[2].text = str(stats["under_min_trips"])
            rc[3].text = str(stats["total_passengers"])

    # ----------------------------------------------------------------
    # Assumptions & Modeling Choices
    # ----------------------------------------------------------------
    doc.add_heading("Assumptions & Modeling Choices", 1)
    doc.add_paragraph(
        "All gaps, inconsistencies, and judgment calls made during modeling are listed below. "
        "These are numbered and also appear in the 'Assumptions' sheet of the workbook."
    )
    for i, a in enumerate(assumptions, 1):
        doc.add_paragraph(f"{i}. {a}", style="List Number")

    # ----------------------------------------------------------------
    # What Changed Since Last Version
    # ----------------------------------------------------------------
    if changelog:
        doc.add_heading("Changes Since Last Version", 1)
        doc.add_paragraph(
            "The following changes were made from the previous model run. "
            "Each entry explains what changed and why."
        )
        for entry in changelog:
            doc.add_paragraph(entry, style="List Bullet")

    doc.save(output_path)
    return output_path
